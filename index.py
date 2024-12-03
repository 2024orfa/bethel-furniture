import os
from dotenv import load_dotenv
import openai
from openai import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pinecone import Pinecone as PineconeClient, ServerlessSpec
import time  
import docx 
# Load environment variables
load_dotenv()

openai_api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=openai_api_key)

# Retrieve environment variables
pinecone_index_name = os.getenv("PINECONE_INDEX_NAME")
pinecone_api_key = os.getenv("PINECONE_API_KEY")
pinecone_environment = os.getenv("PINECONE_ENVIRONMENT")

# Function to extract text from .docx file
def extract_text_from_docx(docx_file_path):
    doc = docx.Document(docx_file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Read the extracted content from the text file (existing code)
file_path = "extracted_content.txt"
with open(file_path, "r", encoding="utf-8") as file:
    text_content = file.read()

# Read and process the .docx file
docx_file_path = "Bethel Furniture FAQ.docx"  # Replace with the actual path to your .docx file
docx_text_content = extract_text_from_docx(docx_file_path)

# Combine the content of both files
combined_text_content = text_content + "\n" + docx_text_content  # Combine .txt and .docx content


# Initialize the RecursiveCharacterTextSplitter
text_splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=150)

# Split the text into chunks
chunks = text_splitter.split_text(combined_text_content)

# Function to get OpenAI embeddings
def get_openai_embeddings(text_list):
    embeddings = []
    batch_size = 5  # Define the batch size based on your needs
    for i in range(0, len(text_list), batch_size):
        batch = text_list[i:i + batch_size]
        try:
            # Ensure correct usage of OpenAI client
            response = client.embeddings.create(
                input=batch,
                model="text-embedding-ada-002"
            )
            # Assuming response.data is a list of embeddings
            batch_embeddings = [item.embedding for item in response.data]  # Use dot notation
            embeddings.extend(batch_embeddings)
            print(f"Processed batch {i // batch_size + 1} of {len(text_list) // batch_size + 1}")
        except openai.error.OpenAIError as e:
            print(f"OpenAI API error for batch starting at index {i}: {e}")
            # Optionally implement retry logic here
    return embeddings

# Embed the chunks using OpenAI
embeddings = get_openai_embeddings(chunks)

# Initialize Pinecone client
client = PineconeClient(api_key=pinecone_api_key, environment=pinecone_environment)

# Check if the index already exists, if not, create it
if pinecone_index_name not in [index.name for index in client.list_indexes()]:
    client.create_index(
        pinecone_index_name,
        dimension=1536,  # The dimension for 'text-embedding-ada-002' is 1536
        metric='cosine',
        spec=ServerlessSpec(
            cloud='aws',
            region='us-east-1'
        )
    )

# Connect to the created index
index = client.Index(pinecone_index_name)

# Prepare and index the data
docs = []
for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
    doc = {
        'id': str(i),
        'values': embedding,  # OpenAI embeddings are already lists
        'metadata': {'text': chunk}
    }
    docs.append(doc)

# Function to split documents into smaller batches
def batch_documents(documents, batch_size):
    for i in range(0, len(documents), batch_size):
        yield documents[i:i + batch_size]

# Define a batch size
batch_size = 100  # Adjust based on Pinecone's upsert limits

# Upsert the documents to Pinecone in batches
for i, batch in enumerate(batch_documents(docs, batch_size), start=1):
    try:
        index.upsert(vectors=batch)
        print(f"Successfully upserted batch {i} of {len(docs) // batch_size + 1}")
        # Optional: Add delay to simulate more visible progress (remove in production)
        time.sleep(0.5)
    except Exception as e:
        print(f"Error during upsert operation for batch {i}: {e}")
        # Optionally implement retry logic here

print(f"Successfully indexed {len(docs)} chunks into Pinecone")


# import os
# from dotenv import load_dotenv
# import openai
# from openai import OpenAI
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from pinecone import Pinecone as PineconeClient, ServerlessSpec
# import time
# import docx
# import pandas as pd
# import json

# # Load environment variables
# load_dotenv()

# openai_api_key = os.getenv("OPENAI_API_KEY")
# client = OpenAI(api_key=openai_api_key)

# pinecone_index_name = os.getenv("PINECONE_INDEX_NAME")
# pinecone_api_key = os.getenv("PINECONE_API_KEY")
# pinecone_environment = os.getenv("PINECONE_ENVIRONMENT")

# # Function to extract text from .docx file
# def extract_text_from_docx(docx_file_path):
#     doc = docx.Document(docx_file_path)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return '\n'.join(full_text)

# # Function to process Excel files and convert them to JSON
# def process_excel_files(file_paths):
#     all_rows = []
#     for file_path in file_paths:
#         try:
#             if file_path.endswith('.xls'):
#                 df = pd.read_excel(file_path, engine='xlrd')  # For .xls files
#             elif file_path.endswith('.xlsx'):
#                 df = pd.read_excel(file_path, engine='openpyxl')  # For .xlsx files
#             else:
#                 print(f"Unsupported file format: {file_path}")
#                 continue

#             json_records = df.to_json(orient='records')
#             records = json.loads(json_records)
#             all_rows.extend(records)
#         except Exception as e:
#             print(f"Error processing {file_path}: {e}")
#     return all_rows

# # Function to sanitize metadata
# def sanitize_metadata(metadata):
#     sanitized_metadata = {}
#     for key, value in metadata.items():
#         if value is None or pd.isnull(value):  # Check for null or NaN values
#             sanitized_metadata[key] = "N/A"  # Replace with a default string
#         elif isinstance(value, list):
#             sanitized_metadata[key] = [str(item) for item in value if item is not None]
#         elif isinstance(value, (str, int, float, bool)):
#             sanitized_metadata[key] = value  # Keep valid types
#         else:
#             sanitized_metadata[key] = str(value)  # Convert unsupported types to strings
#     return sanitized_metadata

# # Function to get OpenAI embeddings
# def get_openai_embeddings(text_list):
#     embeddings = []
#     batch_size = 5
#     for i in range(0, len(text_list), batch_size):
#         batch = text_list[i:i + batch_size]
#         try:
#             response = client.embeddings.create(
#                 input=batch,
#                 model="text-embedding-ada-002"
#             )
#             batch_embeddings = [item.embedding for item in response.data]
#             embeddings.extend(batch_embeddings)
#             print(f"Processed batch {i // batch_size + 1} of {len(text_list) // batch_size + 1}")
#         except openai.error.OpenAIError as e:
#             print(f"OpenAI API error for batch starting at index {i}: {e}")
#     return embeddings

# # Read and process the .docx file
# docx_file_path = "Bethel Furniture FAQ.docx"
# docx_text_content = extract_text_from_docx(docx_file_path)

# # Read the extracted content from the text file
# file_path = "extracted_content.txt"
# with open(file_path, "r", encoding="utf-8") as file:
#     text_content = file.read()

# # Combine text content
# combined_text_content = text_content + "\n" + docx_text_content

# # Initialize RecursiveCharacterTextSplitter
# text_splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=150)
# chunks = text_splitter.split_text(combined_text_content)

# # Embed chunks from text content
# text_embeddings = get_openai_embeddings(chunks)

# # Process Excel files
# excel_file_paths = ["Product list 2.xls", "Product list 3.xls"]  # Add paths to your Excel files
# excel_rows = process_excel_files(excel_file_paths)

# # Convert rows to string for embedding
# row_texts = [json.dumps(row) for row in excel_rows]
# row_embeddings = get_openai_embeddings(row_texts)

# # Combine all embeddings and data
# docs = []
# for i, (chunk, embedding) in enumerate(zip(chunks, text_embeddings)):
#     docs.append({
#         'id': f'text-{i}',
#         'values': embedding,
#         'metadata': {'text': chunk}
#     })

# for i, (row, embedding) in enumerate(zip(excel_rows, row_embeddings)):
#     sanitized_row = sanitize_metadata(row)  # Sanitize the row metadata
#     docs.append({
#         'id': f'row-{i}',
#         'values': embedding,
#         'metadata': sanitized_row
#     })

# # Initialize Pinecone client
# client = PineconeClient(api_key=pinecone_api_key, environment=pinecone_environment)

# # Create index if not exists
# if pinecone_index_name not in [index.name for index in client.list_indexes()]:
#     client.create_index(
#         pinecone_index_name,
#         dimension=1536,
#         metric='cosine',
#         spec=ServerlessSpec(cloud='aws', region='us-east-1')
#     )

# # Connect to the index
# index = client.Index(pinecone_index_name)

# # Function to batch documents
# def batch_documents(documents, batch_size):
#     for i in range(0, len(documents), batch_size):
#         yield documents[i:i + batch_size]

# # Upsert documents into Pinecone
# batch_size = 100
# for i, batch in enumerate(batch_documents(docs, batch_size), start=1):
#     try:
#         index.upsert(vectors=batch)
#         print(f"Successfully upserted batch {i} of {len(docs) // batch_size + 1}")
#         time.sleep(0.5)
#     except Exception as e:
#         print(f"Error during upsert operation for batch {i}: {e}")

# print(f"Successfully indexed {len(docs)} chunks into Pinecone")
